"""Генерация демо-заявок в бинарных форматах (.eml, .docx, .xlsx, .pdf).

Готовые файлы лежат рядом в репозитории; пересобрать их можно одной командой:

    python samples/make_samples.py
"""
from __future__ import annotations

from email.message import EmailMessage
from pathlib import Path

SAMPLES = Path(__file__).resolve().parent


def make_eml() -> None:
    """Письмо с шапкой, подписью и цитатой предыдущего письма."""
    message = EmailMessage()
    message["From"] = "Смирнова Ольга <o.smirnova@agro-trade.ru>"
    message["To"] = "zayavki@rail-operator.ru"
    message["Subject"] = "Заявка на перевозку пшеницы, октябрь"
    message["Date"] = "Mon, 28 Sep 2026 09:14:00 +0300"
    message.set_content(
        "Добрый день!\n\n"
        "ООО «АгроТрейд» просит рассчитать перевозку пшеницы продовольственной\n"
        "со станции Саратов-2 на станцию Новороссийск.\n\n"
        "Объём — 4 200 тонн, это 60 вагонов-зерновозов.\n"
        "Период: с 05.10.2026 по 25.10.2026.\n"
        "Погрузка на элеваторе, две подачи в сутки.\n"
        "Выгрузка в порту, норма выгрузки 4 вагона в час.\n"
        "Ставку просим дать за вагон с НДС.\n\n"
        "С уважением,\n"
        "Смирнова Ольга, руководитель отдела логистики\n"
        "ООО «АгроТрейд», тел. +7 8452 30-40-50\n\n"
        "-----Исходное сообщение-----\n"
        "От: zayavki@rail-operator.ru\n"
        "Тема: Re: сотрудничество\n"
        "Ольга, добрый день! Пришлите, пожалуйста, заявку по форме — "
        "станции, груз, объём, период. Ставку посчитаем в течение дня.\n",
        charset="utf-8",
    )
    (SAMPLES / "04_letter.eml").write_bytes(message.as_bytes())


def make_docx() -> None:
    """Заявка бланком: таблица «поле / значение» плюс текстовые примечания."""
    from docx import Document

    document = Document()
    document.add_heading("ЗАЯВКА НА ЖЕЛЕЗНОДОРОЖНУЮ ПЕРЕВОЗКУ", level=1)
    document.add_paragraph("№ 118 от 29.09.2026")

    rows = [
        ("Заказчик", "ПАО «Северхим»"),
        ("Станция отправления", "Череповец-2"),
        ("Станция назначения", "Усть-Луга"),
        ("Груз", "Удобрения минеральные азотные (карбамид)"),
        ("Род подвижного состава", "Хопперы-минераловозы"),
        ("Количество вагонов", "40"),
        ("Период отгрузки", "01.11.2026 — 30.11.2026"),
        ("Условия погрузки", "Собственный подъездной путь, круглосуточно, норма 3 суток"),
        ("Условия выгрузки", "Терминал порта, выгрузка вагоноопрокидывателем"),
        ("Ставка", "78 000 руб./вагон, без НДС"),
    ]
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = value

    document.add_paragraph("")
    document.add_paragraph(
        "Примечание: возможно увеличение объёма до 55 вагонов при подтверждении "
        "ставки до 25.10.2026."
    )
    document.add_paragraph("Контактное лицо: Гордеев М. В., тел. +7 8202 11-22-33")
    document.save(SAMPLES / "05_zayavka.docx")


def make_xlsx() -> None:
    """План отгрузок таблицей — частый формат заявки от крупного грузоотправителя."""
    from openpyxl import Workbook
    from openpyxl.styles import Font

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "План отгрузок"

    sheet["A1"] = "Заявка на перевозку, АО «Кузбассуголь»"
    sheet["A1"].font = Font(bold=True, size=13)
    sheet["A2"] = "Контакт: Петрова Е. А., e.petrova@kuzbassugol.ru"

    headers = ["Месяц", "Станция отправления", "Станция назначения", "Груз", "Вагонов", "Тонн"]
    for column, title in enumerate(headers, start=1):
        cell = sheet.cell(row=4, column=column, value=title)
        cell.font = Font(bold=True)

    data = [
        ("Ноябрь 2026", "Мыски", "Ванино", "Уголь каменный энергетический", 120, 8280),
        ("Декабрь 2026", "Мыски", "Ванино", "Уголь каменный энергетический", 140, 9660),
    ]
    for row_index, row in enumerate(data, start=5):
        for column, value in enumerate(row, start=1):
            sheet.cell(row=row_index, column=column, value=value)

    sheet["A8"] = "Погрузка: круглосуточно, станция примыкания — Мыски"
    sheet["A9"] = "Выгрузка: порт Ванино, норма 6 суток"
    sheet["A10"] = "Ставка: обсуждается, ориентир 145 000 руб. за вагон без НДС"

    for column, width in zip("ABCDEF", (14, 24, 22, 34, 10, 10)):
        sheet.column_dimensions[column].width = width
    workbook.save(SAMPLES / "06_plan.xlsx")


def make_scan_pdf() -> None:
    """PDF без текстового слоя — имитация скана заявки.

    Текст рисуется линиями, а не шрифтом, поэтому pypdf не извлечёт ни символа:
    ровно та ситуация, в которой система обязана попросить OCR, а не фантазировать.
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    path = SAMPLES / "07_scan.pdf"
    pdf = canvas.Canvas(str(path), pagesize=A4)
    width, height = A4

    pdf.setLineWidth(1.2)
    pdf.rect(40, height - 320, width - 80, 260)
    y = height - 100
    for line_length in (300, 420, 380, 250, 460, 200):
        pdf.line(70, y, 70 + line_length, y)
        y -= 34
    pdf.setLineWidth(3)
    pdf.line(70, height - 340, 260, height - 340)
    pdf.save()


def main() -> None:
    make_eml()
    make_docx()
    make_xlsx()
    make_scan_pdf()
    for name in ("04_letter.eml", "05_zayavka.docx", "06_plan.xlsx", "07_scan.pdf"):
        size = (SAMPLES / name).stat().st_size
        print(f"  создан {name} ({size} байт)")


if __name__ == "__main__":
    main()

"""Word: параграфы плюс таблицы (заявки часто оформлены таблицей «поле / значение»)."""
from __future__ import annotations

import io


def extract(data: bytes) -> tuple[str, list[str]]:
    import docx as python_docx  # импорт внутри, чтобы модуль грузился лениво

    document = python_docx.Document(io.BytesIO(data))
    chunks: list[str] = [p.text.strip() for p in document.paragraphs if p.text.strip()]

    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            # Схлопываем объединённые ячейки, повторяющиеся подряд.
            deduped: list[str] = []
            for cell in cells:
                if not deduped or deduped[-1] != cell:
                    deduped.append(cell)
            line = " | ".join(c for c in deduped if c)
            if line:
                chunks.append(line)

    return "\n".join(chunks), []
